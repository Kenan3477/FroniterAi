"""
Business Document Preprocessor

Comprehensive data preprocessing pipeline for business documents, financial reports,
and regulatory documents for Frontier-1 model training.
"""

import os
import json
import pandas as pd
import numpy as np
from typing import Dict, List, Any, Optional, Tuple, Union
from pathlib import Path
import re
from datetime import datetime, date
import logging
from dataclasses import dataclass
from abc import ABC, abstractmethod
import asyncio
from concurrent.futures import ThreadPoolExecutor
import multiprocessing

# Document processing libraries
import PyPDF2
import docx
from bs4 import BeautifulSoup
import xml.etree.ElementTree as ET

# NLP libraries
import nltk
from nltk.tokenize import sent_tokenize, word_tokenize
from nltk.corpus import stopwords
from nltk.stem import WordNetLemmatizer
import spacy

# ML libraries
from transformers import AutoTokenizer, AutoModel
from sklearn.feature_extraction.text import TfidfVectorizer
from sklearn.model_selection import train_test_split
import torch

# Financial data processing
import yfinance as yf
import pandas_datareader.data as web

logger = logging.getLogger(__name__)

@dataclass
class DocumentMetadata:
    """Metadata for processed documents"""
    source_file: str
    document_type: str
    company_name: Optional[str] = None
    report_period: Optional[str] = None
    filing_date: Optional[date] = None
    industry: Optional[str] = None
    document_length: int = 0
    language: str = "en"
    quality_score: float = 0.0
    
class DocumentType:
    """Document type classifications"""
    FINANCIAL_STATEMENT = "financial_statement"
    ANNUAL_REPORT = "annual_report"
    QUARTERLY_REPORT = "quarterly_report"
    REGULATORY_FILING = "regulatory_filing"
    BUSINESS_PLAN = "business_plan"
    MARKET_RESEARCH = "market_research"
    COMPLIANCE_DOCUMENT = "compliance_document"
    POLICY_DOCUMENT = "policy_document"
    STRATEGIC_PLAN = "strategic_plan"
    AUDIT_REPORT = "audit_report"
    PROSPECTUS = "prospectus"
    PROXY_STATEMENT = "proxy_statement"

@dataclass
class ProcessingConfig:
    """Configuration for document preprocessing"""
    max_document_length: int = 100000
    min_document_length: int = 100
    include_tables: bool = True
    include_images: bool = False
    extract_numbers: bool = True
    normalize_text: bool = True
    remove_stopwords: bool = False
    lemmatize: bool = True
    chunk_size: int = 512
    overlap_size: int = 128
    quality_threshold: float = 0.7
    
class BaseDocumentProcessor(ABC):
    """Base class for document processors"""
    
    def __init__(self, config: ProcessingConfig):
        self.config = config
        self.nlp = spacy.load("en_core_web_sm")
        self.lemmatizer = WordNetLemmatizer()
        
    @abstractmethod
    def extract_text(self, file_path: str) -> str:
        """Extract text from document"""
        pass
    
    @abstractmethod
    def extract_metadata(self, file_path: str, text: str) -> DocumentMetadata:
        """Extract metadata from document"""
        pass
    
    def clean_text(self, text: str) -> str:
        """Clean and normalize text"""
        if not text:
            return ""
        
        # Remove extra whitespace
        text = re.sub(r'\s+', ' ', text)
        
        # Remove special characters but keep financial symbols
        text = re.sub(r'[^\w\s\$\%\(\)\[\]\{\}\-\+\=\.\,\:\;]', ' ', text)
        
        # Normalize financial numbers
        if self.config.extract_numbers:
            text = self._normalize_financial_numbers(text)
        
        # Remove excessive punctuation
        text = re.sub(r'[.,;:]{2,}', '.', text)
        
        return text.strip()
    
    def _normalize_financial_numbers(self, text: str) -> str:
        """Normalize financial numbers and currencies"""
        # Convert currency formats
        text = re.sub(r'\$(\d+(?:,\d{3})*(?:\.\d{2})?)', r'USD_\1', text)
        text = re.sub(r'(\d+(?:,\d{3})*(?:\.\d{2})?)\s*million', r'\1M', text)
        text = re.sub(r'(\d+(?:,\d{3})*(?:\.\d{2})?)\s*billion', r'\1B', text)
        text = re.sub(r'(\d+(?:,\d{3})*(?:\.\d{2})?)\s*trillion', r'\1T', text)
        
        # Normalize percentages
        text = re.sub(r'(\d+(?:\.\d+)?)\s*percent', r'\1%', text)
        
        return text
    
    def calculate_quality_score(self, text: str, metadata: DocumentMetadata) -> float:
        """Calculate document quality score"""
        score = 0.0
        
        # Length score (normalized)
        length_score = min(len(text) / self.config.max_document_length, 1.0)
        score += length_score * 0.3
        
        # Readability score
        sentences = sent_tokenize(text)
        if sentences:
            avg_sentence_length = len(word_tokenize(text)) / len(sentences)
            readability_score = 1.0 / (1.0 + abs(avg_sentence_length - 20) / 20)
            score += readability_score * 0.2
        
        # Financial content score
        financial_terms = [
            'revenue', 'profit', 'loss', 'assets', 'liabilities', 'equity',
            'cash flow', 'ebitda', 'margin', 'ratio', 'financial', 'statement'
        ]
        financial_score = sum(1 for term in financial_terms if term in text.lower())
        financial_score = min(financial_score / len(financial_terms), 1.0)
        score += financial_score * 0.3
        
        # Structure score (presence of sections, tables, etc.)
        structure_indicators = ['table', 'section', 'chapter', 'exhibit', 'appendix']
        structure_score = sum(1 for indicator in structure_indicators if indicator in text.lower())
        structure_score = min(structure_score / len(structure_indicators), 1.0)
        score += structure_score * 0.2
        
        return min(score, 1.0)

class PDFProcessor(BaseDocumentProcessor):
    """Processor for PDF documents"""
    
    def extract_text(self, file_path: str) -> str:
        """Extract text from PDF"""
        try:
            text = ""
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                for page in pdf_reader.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text += page_text + "\n"
            
            return self.clean_text(text)
            
        except Exception as e:
            logger.error(f"Error extracting text from PDF {file_path}: {e}")
            return ""
    
    def extract_metadata(self, file_path: str, text: str) -> DocumentMetadata:
        """Extract metadata from PDF"""
        metadata = DocumentMetadata(
            source_file=file_path,
            document_type=self._detect_document_type(text),
            document_length=len(text)
        )
        
        # Extract company name
        metadata.company_name = self._extract_company_name(text)
        
        # Extract report period
        metadata.report_period = self._extract_report_period(text)
        
        # Extract filing date
        metadata.filing_date = self._extract_filing_date(text)
        
        # Calculate quality score
        metadata.quality_score = self.calculate_quality_score(text, metadata)
        
        return metadata
    
    def _detect_document_type(self, text: str) -> str:
        """Detect document type from content"""
        text_lower = text.lower()
        
        if any(term in text_lower for term in ['10-k', '10-q', '8-k', 'sec filing']):
            if '10-k' in text_lower:
                return DocumentType.ANNUAL_REPORT
            elif '10-q' in text_lower:
                return DocumentType.QUARTERLY_REPORT
            else:
                return DocumentType.REGULATORY_FILING
        
        if any(term in text_lower for term in ['financial statement', 'balance sheet', 'income statement']):
            return DocumentType.FINANCIAL_STATEMENT
        
        if any(term in text_lower for term in ['annual report', 'yearly report']):
            return DocumentType.ANNUAL_REPORT
        
        if any(term in text_lower for term in ['business plan', 'strategic plan']):
            return DocumentType.BUSINESS_PLAN
        
        if any(term in text_lower for term in ['audit report', 'auditor']):
            return DocumentType.AUDIT_REPORT
        
        return DocumentType.BUSINESS_PLAN  # Default
    
    def _extract_company_name(self, text: str) -> Optional[str]:
        """Extract company name from document"""
        # Look for company name patterns
        patterns = [
            r'([A-Z][a-z]+(?:\s+[A-Z][a-z]+)*)\s+(?:Inc|Corp|Corporation|LLC|Ltd|Limited)',
            r'Company:\s*([A-Z][a-zA-Z\s&]+)',
            r'([A-Z][a-zA-Z\s&]+)\s+Financial\s+Statements?'
        ]
        
        for pattern in patterns:
            match = re.search(pattern, text)
            if match:
                return match.group(1).strip()
        
        return None
    
    def _extract_report_period(self, text: str) -> Optional[str]:
        """Extract reporting period from document"""
        patterns = [
            r'(?:Quarter|Period)\s+(?:Ended|Ending)\s+([A-Z][a-z]+\s+\d{1,2},?\s+\d{4})',
            r'(?:Year|Fiscal\s+Year)\s+(?:Ended|Ending)\s+([A-Z][a-z]+\s+\d{1,2},?\s+\d{4})',
            r'For\s+the\s+(?:quarter|year)\s+ended\s+([A-Z][a-z]+\s+\d{1,2},?\s+\d{4})',
            r'Q[1-4]\s+(\d{4})',
            r'(\d{4})\s+Annual\s+Report'
        ]
        
        for pattern in patterns:
            match = re.search(pattern, text, re.IGNORECASE)
            if match:
                return match.group(1).strip()
        
        return None
    
    def _extract_filing_date(self, text: str) -> Optional[date]:
        """Extract filing date from document"""
        patterns = [
            r'Filed?\s+on:?\s*([A-Z][a-z]+\s+\d{1,2},?\s+\d{4})',
            r'Filing\s+Date:?\s*([A-Z][a-z]+\s+\d{1,2},?\s+\d{4})',
            r'Date\s+Filed:?\s*([A-Z][a-z]+\s+\d{1,2},?\s+\d{4})'
        ]
        
        for pattern in patterns:
            match = re.search(pattern, text, re.IGNORECASE)
            if match:
                try:
                    date_str = match.group(1).strip()
                    return datetime.strptime(date_str, "%B %d, %Y").date()
                except ValueError:
                    continue
        
        return None

class ExcelProcessor(BaseDocumentProcessor):
    """Processor for Excel financial documents"""
    
    def extract_text(self, file_path: str) -> str:
        """Extract text from Excel file"""
        try:
            # Read all sheets
            xl_file = pd.ExcelFile(file_path)
            text_parts = []
            
            for sheet_name in xl_file.sheet_names:
                df = pd.read_excel(file_path, sheet_name=sheet_name)
                
                # Convert dataframe to text
                sheet_text = f"Sheet: {sheet_name}\n"
                sheet_text += df.to_string(index=False, na_rep='')
                text_parts.append(sheet_text)
            
            return self.clean_text("\n\n".join(text_parts))
            
        except Exception as e:
            logger.error(f"Error extracting text from Excel {file_path}: {e}")
            return ""
    
    def extract_metadata(self, file_path: str, text: str) -> DocumentMetadata:
        """Extract metadata from Excel file"""
        metadata = DocumentMetadata(
            source_file=file_path,
            document_type=DocumentType.FINANCIAL_STATEMENT,
            document_length=len(text)
        )
        
        # Excel files are typically financial statements
        metadata.document_type = DocumentType.FINANCIAL_STATEMENT
        metadata.quality_score = self.calculate_quality_score(text, metadata)
        
        return metadata

class WordProcessor(BaseDocumentProcessor):
    """Processor for Word documents"""
    
    def extract_text(self, file_path: str) -> str:
        """Extract text from Word document"""
        try:
            doc = docx.Document(file_path)
            text_parts = []
            
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_parts.append(paragraph.text)
            
            # Extract table data if configured
            if self.config.include_tables:
                for table in doc.tables:
                    table_text = []
                    for row in table.rows:
                        row_text = []
                        for cell in row.cells:
                            if cell.text.strip():
                                row_text.append(cell.text.strip())
                        if row_text:
                            table_text.append(" | ".join(row_text))
                    if table_text:
                        text_parts.append("\n".join(table_text))
            
            return self.clean_text("\n".join(text_parts))
            
        except Exception as e:
            logger.error(f"Error extracting text from Word {file_path}: {e}")
            return ""
    
    def extract_metadata(self, file_path: str, text: str) -> DocumentMetadata:
        """Extract metadata from Word document"""
        metadata = DocumentMetadata(
            source_file=file_path,
            document_type=self._detect_document_type(text),
            document_length=len(text)
        )
        
        metadata.quality_score = self.calculate_quality_score(text, metadata)
        return metadata

class BusinessDocumentPreprocessor:
    """Main preprocessor for business documents"""
    
    def __init__(self, config: ProcessingConfig = None):
        self.config = config or ProcessingConfig()
        self.processors = {
            '.pdf': PDFProcessor(self.config),
            '.docx': WordProcessor(self.config),
            '.doc': WordProcessor(self.config),
            '.xlsx': ExcelProcessor(self.config),
            '.xls': ExcelProcessor(self.config)
        }
        
        # Initialize tokenizer for chunking
        self.tokenizer = AutoTokenizer.from_pretrained("microsoft/DialoGPT-medium")
        
        # Download required NLTK data
        try:
            nltk.download('punkt', quiet=True)
            nltk.download('stopwords', quiet=True)
            nltk.download('wordnet', quiet=True)
        except:
            pass
    
    def process_directory(self, input_dir: str, output_dir: str) -> Dict[str, Any]:
        """Process all documents in a directory"""
        input_path = Path(input_dir)
        output_path = Path(output_dir)
        output_path.mkdir(parents=True, exist_ok=True)
        
        results = {
            'processed_files': [],
            'failed_files': [],
            'total_documents': 0,
            'total_chunks': 0,
            'quality_stats': {
                'high_quality': 0,
                'medium_quality': 0,
                'low_quality': 0
            }
        }
        
        # Get all supported files
        supported_extensions = list(self.processors.keys())
        files = []
        for ext in supported_extensions:
            files.extend(input_path.glob(f"**/*{ext}"))
        
        logger.info(f"Found {len(files)} documents to process")
        
        # Process files with multiprocessing
        with ThreadPoolExecutor(max_workers=multiprocessing.cpu_count()) as executor:
            futures = []
            for file_path in files:
                future = executor.submit(self._process_single_file, file_path, output_path)
                futures.append((file_path, future))
            
            for file_path, future in futures:
                try:
                    result = future.result()
                    if result:
                        results['processed_files'].append(str(file_path))
                        results['total_documents'] += 1
                        results['total_chunks'] += result['chunk_count']
                        
                        # Update quality stats
                        quality = result['metadata']['quality_score']
                        if quality >= 0.8:
                            results['quality_stats']['high_quality'] += 1
                        elif quality >= 0.6:
                            results['quality_stats']['medium_quality'] += 1
                        else:
                            results['quality_stats']['low_quality'] += 1
                    else:
                        results['failed_files'].append(str(file_path))
                        
                except Exception as e:
                    logger.error(f"Error processing {file_path}: {e}")
                    results['failed_files'].append(str(file_path))
        
        # Save processing summary
        summary_file = output_path / "processing_summary.json"
        with open(summary_file, 'w') as f:
            json.dump(results, f, indent=2, default=str)
        
        logger.info(f"Processing complete. Summary saved to {summary_file}")
        return results
    
    def _process_single_file(self, file_path: Path, output_dir: Path) -> Optional[Dict[str, Any]]:
        """Process a single file"""
        try:
            # Get appropriate processor
            processor = self.processors.get(file_path.suffix.lower())
            if not processor:
                logger.warning(f"No processor for file type: {file_path.suffix}")
                return None
            
            # Extract text and metadata
            text = processor.extract_text(str(file_path))
            if not text or len(text) < self.config.min_document_length:
                logger.warning(f"Document too short or empty: {file_path}")
                return None
            
            metadata = processor.extract_metadata(str(file_path), text)
            
            # Check quality threshold
            if metadata.quality_score < self.config.quality_threshold:
                logger.warning(f"Document quality too low: {file_path} (score: {metadata.quality_score})")
                return None
            
            # Create chunks
            chunks = self._create_chunks(text)
            
            # Save processed data
            output_file = output_dir / f"{file_path.stem}.json"
            processed_data = {
                'metadata': {
                    'source_file': str(file_path),
                    'document_type': metadata.document_type,
                    'company_name': metadata.company_name,
                    'report_period': metadata.report_period,
                    'filing_date': metadata.filing_date,
                    'document_length': metadata.document_length,
                    'quality_score': metadata.quality_score,
                    'chunk_count': len(chunks),
                    'processed_at': datetime.now().isoformat()
                },
                'chunks': chunks
            }
            
            with open(output_file, 'w', encoding='utf-8') as f:
                json.dump(processed_data, f, indent=2, default=str)
            
            return {
                'chunk_count': len(chunks),
                'metadata': processed_data['metadata']
            }
            
        except Exception as e:
            logger.error(f"Error processing file {file_path}: {e}")
            return None
    
    def _create_chunks(self, text: str) -> List[Dict[str, Any]]:
        """Create overlapping chunks from text"""
        # Tokenize text
        tokens = self.tokenizer.encode(text, add_special_tokens=False)
        
        chunks = []
        chunk_id = 0
        
        for i in range(0, len(tokens), self.config.chunk_size - self.config.overlap_size):
            chunk_tokens = tokens[i:i + self.config.chunk_size]
            
            if len(chunk_tokens) < 50:  # Skip very small chunks
                continue
            
            chunk_text = self.tokenizer.decode(chunk_tokens, skip_special_tokens=True)
            
            chunk_data = {
                'chunk_id': chunk_id,
                'text': chunk_text,
                'token_count': len(chunk_tokens),
                'start_position': i,
                'end_position': i + len(chunk_tokens)
            }
            
            chunks.append(chunk_data)
            chunk_id += 1
        
        return chunks
    
    def create_training_dataset(self, processed_dir: str, output_file: str, test_split: float = 0.2):
        """Create training dataset from processed documents"""
        processed_path = Path(processed_dir)
        
        all_chunks = []
        all_metadata = []
        
        # Collect all chunks
        for json_file in processed_path.glob("*.json"):
            try:
                with open(json_file, 'r', encoding='utf-8') as f:
                    data = json.load(f)
                
                for chunk in data['chunks']:
                    all_chunks.append({
                        'text': chunk['text'],
                        'document_type': data['metadata']['document_type'],
                        'company_name': data['metadata']['company_name'],
                        'quality_score': data['metadata']['quality_score'],
                        'token_count': chunk['token_count']
                    })
                
                all_metadata.append(data['metadata'])
                
            except Exception as e:
                logger.error(f"Error reading {json_file}: {e}")
        
        logger.info(f"Collected {len(all_chunks)} chunks from {len(all_metadata)} documents")
        
        # Split into train/test
        train_chunks, test_chunks = train_test_split(
            all_chunks, 
            test_size=test_split, 
            random_state=42,
            stratify=[chunk['document_type'] for chunk in all_chunks]
        )
        
        # Save dataset
        dataset = {
            'train': train_chunks,
            'test': test_chunks,
            'metadata': {
                'total_chunks': len(all_chunks),
                'train_size': len(train_chunks),
                'test_size': len(test_chunks),
                'document_types': list(set(chunk['document_type'] for chunk in all_chunks)),
                'created_at': datetime.now().isoformat()
            }
        }
        
        with open(output_file, 'w', encoding='utf-8') as f:
            json.dump(dataset, f, indent=2, default=str)
        
        logger.info(f"Training dataset saved to {output_file}")
        return dataset

class FinancialDataProcessor:
    """Processor for financial market data"""
    
    def __init__(self):
        self.companies = []
        
    def fetch_market_data(self, symbols: List[str], period: str = "5y") -> pd.DataFrame:
        """Fetch market data for training"""
        try:
            data = yf.download(symbols, period=period)
            return data
        except Exception as e:
            logger.error(f"Error fetching market data: {e}")
            return pd.DataFrame()
    
    def create_financial_features(self, data: pd.DataFrame) -> pd.DataFrame:
        """Create financial features for training"""
        features = data.copy()
        
        # Technical indicators
        features['SMA_20'] = data['Close'].rolling(window=20).mean()
        features['SMA_50'] = data['Close'].rolling(window=50).mean()
        features['RSI'] = self._calculate_rsi(data['Close'])
        features['Volatility'] = data['Close'].rolling(window=20).std()
        
        # Price changes
        features['Daily_Return'] = data['Close'].pct_change()
        features['Weekly_Return'] = data['Close'].pct_change(periods=5)
        features['Monthly_Return'] = data['Close'].pct_change(periods=21)
        
        return features
    
    def _calculate_rsi(self, prices: pd.Series, window: int = 14) -> pd.Series:
        """Calculate RSI indicator"""
        delta = prices.diff()
        gain = (delta.where(delta > 0, 0)).rolling(window=window).mean()
        loss = (-delta.where(delta < 0, 0)).rolling(window=window).mean()
        rs = gain / loss
        rsi = 100 - (100 / (1 + rs))
        return rsi

# Example usage and main processing functions
def main():
    """Main preprocessing pipeline"""
    
    # Configuration
    config = ProcessingConfig(
        max_document_length=50000,
        min_document_length=500,
        include_tables=True,
        extract_numbers=True,
        normalize_text=True,
        chunk_size=512,
        overlap_size=128,
        quality_threshold=0.6
    )
    
    # Initialize preprocessor
    preprocessor = BusinessDocumentPreprocessor(config)
    
    # Process documents
    results = preprocessor.process_directory(
        input_dir="./raw_documents",
        output_dir="./processed_documents"
    )
    
    # Create training dataset
    dataset = preprocessor.create_training_dataset(
        processed_dir="./processed_documents",
        output_file="./training_dataset.json",
        test_split=0.2
    )
    
    print(f"Processing complete:")
    print(f"- Processed {results['total_documents']} documents")
    print(f"- Created {results['total_chunks']} chunks")
    print(f"- Training dataset: {dataset['metadata']['train_size']} chunks")
    print(f"- Test dataset: {dataset['metadata']['test_size']} chunks")

if __name__ == "__main__":
    logging.basicConfig(level=logging.INFO)
    main()
